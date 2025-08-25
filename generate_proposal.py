import pandas as pd
from openai import OpenAI
from docx import Document
from datetime import datetime
import json
import os
import re

# Today's Date:
today_date = datetime.today()
today = today_date.strftime('%m/%d/%y')


def Load_Templates(BID_TEMPLATE_FILE, SOl_INFO_TEMPLATE):
    # Load the template text file
    with open(BID_TEMPLATE_FILE, "r", encoding="utf-8") as f:
        template_text = f.read()

    with open(SOl_INFO_TEMPLATE, "r", encoding="utf-8") as f:
        sol_info_text = f.read()
    return template_text, sol_info_text

# Load CSV
def load_csv(csv_file):
    df = pd.read_csv(csv_file)
    return df

# Function to call GPT for summarizing description and extracting instructions
def generate_prompt(description, instructions, listed_supplier, listed_part, supplier_options, quantity):
    part_number_prompt = f"The recommended part number for this solicitation is {listed_part}."
    if listed_part.lower() == "not listed" or listed_part.lower() == "not specified" or listed_part.strip().lower() == "nan":
        part_number_prompt = "The solicitation did not specify a part number or sku for this item."

    suggested_supplier_prompt = f"The recommended manufacturer or supplier for this solicitation is {listed_supplier}."
    if listed_supplier.lower() == "not listed" or listed_supplier.lower() == "not specified" or listed_supplier.strip().lower() == "nan":
        suggested_supplier_prompt = "The solicitation did not specify a manufacturer or supplier for this item."

    prompt =  f"""
        You are a government proposal assistant. Given the following product information obtained from a government solicitation and list of potential suppliers, determine the best supplier and help gather information
        required to submit a bid/proposal for the solicitation.


        Description of the product obtained from the solicitation:

        {description}
        ---
        The solicitation asks for a quantity of {quantity} of the specified item.
        ---
        {part_number_prompt}
        ---
        {suggested_supplier_prompt}
        ---
        Submission Instructions from Solicitation:

        {instructions}

        
        Optional Supplier 1 Details:
        {supplier_options[0]}

        Optional Supplier 2 Details:
        {supplier_options[1]}

        Optional Supplier 3 Details:
        {supplier_options[2]}

        Based on this information, determine the best supplier by factoring in unit price (looking for lowest price) and match of product and supplier with solicitation requirements/suggestions.
        Ensure suppliers actually supply parts and are not companies that just post government solicitations.

        Provide the following information
        - Recommended supplier (Supplier 1, Supplier 2, or Supplier 3)
        - shipping for government entity or N/A if not available.
        - reasoning for recommendation decision.

        Respond ONLY with raw JSON — do not wrap it in quotes or formatting code blocks. Return valid JSON with double-escaped backslashes and escaped internal quotes where needed.
        Do not include any additional commentary.

        Example format:

        [
          {{
            "Recommended Supplier": "Supplier 2",
            "shipping address": "1234 Walnut Ave. Newark, NJ 12345",
            "Reasoning": "The part from supplier 2 is most similar to the part in the solicitation."
          }}
        ]
        If no appropriate supplier is found from the supplier information above return "N/A" in both "Recommended Supplier" and "shipping address" categories in the JSON output.
        Do not search for additional suppliers. Only use suppliers above.
    """
    return prompt

def extract_json(text):
    # Remove optional markdown code fences if present
    return re.sub(r"^```(?:json)?|```$", "", text.strip(), flags=re.MULTILINE).strip()

# Process each valid quote
def validate_supplier_and_write_proposal(df, output_directory, Open_AI_API_Key, BID_TEMPLATE_FILE, SOl_INFO_TEMPLATE):
    client = OpenAI(api_key=Open_AI_API_Key)
    for index, row in df.iterrows():
        try:
            solicitation_number = row["solicitation number"]
            notice_id = row["notice id"]
            notice_type = row["notice type"]
            NAICS_code = row["NAICS Code"]
            posted_date = row["posted date"]
            due_date = row["due date"]
            description = row["item description"]
            title = row["title"]
            instructions = row["submission instructions"]
            fulfillment = row["fulfillment instructions"]
            quantity = row["quantity"]
            poc_name = row["poc name"]
            poc_email = row["poc email"]
            listed_supplier = row["listed supplier"]
            listed_part_number = row["item part#"]
            link = row["solicitation link"]

            supplier_1 = row["supplier 1"]
            supplier_1_rank = row["supplier 1 rank"]
            supplier_1_part_number = row["supplier 1 product #"]
            supplier_1_part_name = row["supplier 1 product name"]
            price_1 = row["supplier 1 quote"]
            supplier_1_link = row["supplier 1 link"]
            supplier_2 = row["supplier 2"]
            supplier_2_rank = row["supplier 2 rank"]
            supplier_2_part_number = row["supplier 2 product #"]
            supplier_2_part_name = row["supplier 2 product name"]
            price_2 = row["supplier 2 quote"]
            supplier_2_link = row["supplier 2 link"]
            supplier_3 = row["supplier 3"]
            supplier_3_rank = row["supplier 3 rank"]
            supplier_3_part_number = row["supplier 3 product #"]
            supplier_3_part_name = row["supplier 3 product name"]
            price_3 = row["supplier 3 quote"]
            supplier_3_link = row["supplier 3 link"]

            suppliers =[
                {
                "supplier": str(supplier_1),
                "product #": str(supplier_1_part_number),
                "product name": str(supplier_1_part_name),
                "unit price": str(price_1),
                },
                {
                "supplier": str(supplier_2),
                "product #": str(supplier_2_part_number),
                "product name": str(supplier_2_part_name),
                "unit price": str(price_2),
                },
                {
                "supplier": str(supplier_3),
                "product #": str(supplier_3_part_number),
                "product name": str(supplier_3_part_name),
                "unit price": str(price_3),
                }
            ]

            #make directory
            os.makedirs(f"{output_directory}/{solicitation_number}", exist_ok=True)
            
            template_text, sol_info_text = Load_Templates(BID_TEMPLATE_FILE, SOl_INFO_TEMPLATE)

            # Fill in solicitation details template
            filled_text = sol_info_text
            filled_text = filled_text.replace("{posted_date}", str(posted_date))
            filled_text = filled_text.replace("{due_date}", str(due_date))
            filled_text = filled_text.replace("{title}", str(title))
            filled_text = filled_text.replace("{notice_id}", str(notice_id))
            filled_text = filled_text.replace("{notice_type}", str(notice_type))
            filled_text = filled_text.replace("{solicitation_number}", str(solicitation_number))
            filled_text = filled_text.replace("{NAICS_code}", str(NAICS_code))
            filled_text = filled_text.replace("{description}", str(description))
            filled_text = filled_text.replace("{suggested_part_number}", str(listed_part_number))
            filled_text = filled_text.replace("{quantity}", str(quantity))
            filled_text = filled_text.replace("{suggested_supplier}", str(listed_supplier))
            filled_text = filled_text.replace("{poc_name}", str(poc_name))
            filled_text = filled_text.replace("{poc_email}", str(poc_email))
            filled_text = filled_text.replace("{instructions}", str(instructions))
            filled_text = filled_text.replace("{fulfillment_instructions}", str(fulfillment))
            filled_text = filled_text.replace("{link}", str(link))
            filled_text = filled_text.replace("{supplier_1}", str(supplier_1))
            filled_text = filled_text.replace("{supplier_1_rank}", str(supplier_1_rank))
            filled_text = filled_text.replace("{supplier_1_product_number}", str(supplier_1_part_number))
            filled_text = filled_text.replace("{supplier_1_product_name}", str(supplier_1_part_name))
            filled_text = filled_text.replace("{supplier_1_unit_price}", str(price_1))
            filled_text = filled_text.replace("{supplier_1_link}", str(supplier_1_link))
            filled_text = filled_text.replace("{supplier_2}", str(supplier_2))
            filled_text = filled_text.replace("{supplier_2_rank}", str(supplier_2_rank))
            filled_text = filled_text.replace("{supplier_2_product_number}", str(supplier_2_part_number))
            filled_text = filled_text.replace("{supplier_2_product_name}", str(supplier_2_part_name))
            filled_text = filled_text.replace("{supplier_2_unit_price}", str(price_2))
            filled_text = filled_text.replace("{supplier_2_link}", str(supplier_2_link))
            filled_text = filled_text.replace("{supplier_3}", str(supplier_3))
            filled_text = filled_text.replace("{supplier_3_rank}", str(supplier_3_rank))
            filled_text = filled_text.replace("{supplier_3_product_number}", str(supplier_3_part_number))
            filled_text = filled_text.replace("{supplier_3_product_name}", str(supplier_3_part_name))
            filled_text = filled_text.replace("{supplier_3_unit_price}", str(price_3))
            filled_text = filled_text.replace("{supplier_3_link}", str(supplier_3_link))
            

            # Convert to Word Document
            doc = Document()
            for line in filled_text.split("\n"):
                doc.add_paragraph(line)

            filename = f"{output_directory}/{solicitation_number}/{today.replace("/", "")}_{solicitation_number}_INFO.docx"
            doc.save(filename)
                        
            for supplier in suppliers:
                supplier_name = supplier.get("supplier")
                product_number = supplier.get("product #")
                product_name = supplier.get("product name")
                unit_price = supplier.get("unit price")

                # Fill in proposal template
                filled_text = template_text
                filled_text = filled_text.replace("{solicitation_number}", str(solicitation_number))
                if product_name.strip().lower() != "nan":
                    filled_text = filled_text.replace("{product_title}", str(product_name))
                if product_number.strip().lower() != "nan":
                    filled_text = filled_text.replace("{product_number}", str(product_number))
                filled_text = filled_text.replace("{description}", str(description))
                filled_text = filled_text.replace("{quantity}", str(quantity))
                if unit_price.strip().lower() != "nan":
                    filled_text = filled_text.replace("{price_ea}", f"${unit_price}")
                filled_text = filled_text.replace("{Today}", str(today))
                

                # Convert to Word Document
                doc = Document()
                for line in filled_text.split("\n"):
                    doc.add_paragraph(line)
                if str(supplier_name) == "N/A":
                    supplier_name = "None"   
                filename = f"{output_directory}/{solicitation_number}/{today.replace("/", "")}_{solicitation_number}_{supplier_name}.docx"
                doc.save(filename)

        except Exception as e:
            print(f"⚠️ Error processing row {index}: {e}")

#Run Script
if __name__ == "__main__":

    # === File paths ===
    CSV_FILE = "supplier_suggestions.csv"
    BID_TEMPLATE_FILE = "bid_template.txt"
    SOl_INFO_TEMPLATE = "solicitation_details_template.txt"
    OUTPUT_DIR = "proposals"

    OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")  # use this instead of a literal "sk-..."    
    df = load_csv(CSV_FILE)
    validate_supplier_and_write_proposal(df, OUTPUT_DIR, OPENAI_API_KEY, BID_TEMPLATE_FILE, SOl_INFO_TEMPLATE)
